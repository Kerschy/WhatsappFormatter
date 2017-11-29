import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

path = "assets/"
regexp = re.compile(r'^(0[1-9]|[12][0-9]|3[01])\.(0[1-9]|1[012])\.\d\d')

def formatFile(file):
    newContent = ""
    fullLine = ""
    hasContent = False
    for line in file:
        if (regexp.match(line)):
            if(hasContent):
                hasContent = False
            newContent += fullLine
            fullLine = line
        else:
            fullLine = fullLine.replace("\n", " ")
            fullLine += line
            hasContent = True
    file.seek(0)
    file.truncate()
    file.write(newContent)


def transformFile(file):
    document = Document()
    oldDate = ""
    lineNumber = 1
    for line in file:
        date = line.split(",",1)[0]
        date = date.split(".")
        date[2] = "20" + date[2]
        date = date[0] + "." + date[1] + "." + date[2]

        if (date != oldDate):
            p = document.add_paragraph()
            run = p.add_run("\n" + date)
            run.bold = True
            font = run.font
            font.size = Pt(12)
            paragraph_format = p.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            oldDate = date

        t = document.add_paragraph()
        time = line.split(", ",1)[1]
        timeNoSec = time.split(":", 2)
        timeNoSec = timeNoSec[0] + ":" + timeNoSec[1]
        t.add_run(timeNoSec).italic = True

        if(line.count(':') > 3):
            name = time.split(":")[3] + ":\t"
            t.add_run(name).bold = True
            text = time.split(":")[4].replace(" ", "", 1).replace("\n", " ")
            t.add_run(text)
        else:
            text = time.split(":", 3)[3].replace(": ", ":\t", 1).replace("\n", " ")
            t.add_run(text)

        tFormat = t.paragraph_format
        #print(text)
        tFormat.left_indent = Inches(2.5)
        tFormat.first_line_indent = Inches(-2.5)
        print(lineNumber)
        lineNumber += 1
    documentName = (file.name).replace(".txt", ".docx")
    document.save(documentName)

for filename in os.listdir(path):
    if filename.endswith(".txt"):
        filename = os.path.join(path, filename)
        file = open(filename, 'r+')
        formatFile(file)
        file.close()
        file = open(filename, 'r+')
        transformFile(file)
        file.close()